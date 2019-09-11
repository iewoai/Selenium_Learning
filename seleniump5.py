from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.common.action_chains import ActionChains
import chardet, time, random
import re, os
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from tqdm import tqdm

file = 'gdCMAe-Part2-10.docx'
document = Document()
# 全局调整正文字体样式、大小(全局调整导致图片格式出错)
document.styles['Normal'].font.name='宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
document.styles['Normal'].font.size = Pt(10.5)


url = 'https://tiku.gaodun.com/nexam/newcmapart2'
file_path = 'F:\\py学习\\selenium\\gdCMA'
def is_element_exist(xpath):
	try:
		time.sleep(0.2)
		browser.find_element_by_xpath(xpath)
		return True
	except:
		return False
t0 = time.time()
p_ids = ['10443']#,

for pid_set,pid in enumerate(p_ids):
	browser = webdriver.Chrome()

	browser.get(url)
	# browser.maximize_window()

	time.sleep(2)
	browser.find_element_by_xpath('//em[@class="tongjiClk"]').click()
	time.sleep(1)
	browser.find_element_by_xpath('//div[@class="tab-item"]/a').click()

	browser.find_element_by_xpath('//input[@id="user"]').send_keys('15757933663')
	browser.find_element_by_xpath('//input[@id="pass"]').send_keys('ZQJpqtyxz8426.')

	browser.find_element_by_xpath('//button[@class="login-btn"]').click()
	time.sleep(1)

	# browser.find_element_by_xpath('//ul[@class="nav"]/li[3]/a//i[@class="fa fa-pluswbk"]').click()
	# time.sleep(0.5)
	# browser.find_element_by_xpath('//ul[@class="nav"]/li[3]/ul[@class="ch-ulOne"]/li[1]/a//i[@class="fa fa-pluswbk"]').click()
	# time.sleep(1)
	# title = browser.find_element_by_xpath('//a[@type_val="%s"]/parent::p/parent::li[@class="ch-curMdul"]' % pid).get_attribute('papercategory')
	browser.find_element_by_xpath('//a[@type_val="%s"]' % pid).click()
	time.sleep(1)
	browser.find_elements_by_xpath('//div[@class="col-lg-3 text-center pa-pattBox"]')[1].click()

	wait = WebDriverWait(browser, 10)
	wait.until(EC.presence_of_element_located((By.XPATH, '//em[@data-value="A"]')))

	next_qs = browser.find_elements_by_xpath('//div[@class="f-doExam-lst"]')
	# print(len(next_qs))

	# p = document.add_paragraph()
	# p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	# run = p.add_run(title)
	# run.bold = True
	# run.font.size = Pt(16)

	# paragraph = document.add_paragraph('Cost Management')
	# paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



	# document.add_paragraph('Internal Controls').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	for i,next_q in tqdm(enumerate(next_qs)):
		# print('进度：%d/%d' % (i, len(next_qs)-1))
		# print(next_q.is_displayed())
		# 题目id
		q_id = next_q.get_attribute('id')
		q_id = re.findall('Id_(\d+)', q_id)
		q_id = q_id[0]
		# 题目号
		q_type = browser.find_elements_by_xpath('//div[@class="type"]/em')[i].get_attribute('textContent')
		q_type = q_type.encode('GBK','ignore').decode('GBk')

		document.add_paragraph(q_type).paragraph_format.line_spacing = Pt(22)
		# print(q_type)
		# 问题
		stem_xpath = '//td[@id="keyWord_%s"]' % q_id
		q_stem = browser.find_element_by_xpath(stem_xpath).get_attribute('textContent')

		document.add_paragraph(q_stem).paragraph_format.line_spacing = Pt(22)
		# print(q_stem)
		# 获取题目中可能出现的配图
		img_xpath = '//td[@id="keyWord_%s"]//img' % q_id
		if is_element_exist(img_xpath):
			
			img = browser.find_elements_by_xpath(img_xpath)
			for j,im in enumerate(img):
				imgname = q_id + '_stem-%d.png' % j
				img_path = os.path.join(file_path, imgname)
				# print('stem图片地址为：%s' % (img_path, ))
				img_src = im.get_attribute('src')
				try:
					if not os.path.exists(img_path):
						img = requests.get(img_src)
						with open(img_path, 'wb') as f:
							f.write(img.content)
					pic = document.add_picture(img_path)
					last_paragraph = document.paragraphs[-1]
					last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
				except:
					pass
		# 选择
		q_opt = browser.find_elements_by_xpath('//div[@class="sep-opt"]/table/tbody/tr/td')[i].get_attribute('textContent')
		q_opt = q_opt.encode('GBK','ignore').decode('GBk')
		q_opt = re.sub('([A|B|C|D])\.', r'\n\1.', q_opt)
		q_opt = re.sub('([A|B|C|D]、)', r'\n\1', q_opt)

		document.add_paragraph(q_opt).paragraph_format.line_spacing = Pt(22)
		# print(q_opt)
		# 判断正确答案不可见时，点击A选项
		# 正确答案
		# print('正确答案为：')
		q_right = browser.find_elements_by_xpath('//em[@class="d-right"]')[i]

		# print(q_right.is_displayed())
		# if not q_right.is_displayed():
		# 	browser.find_elements_by_xpath('//em[@data-value="A"]/span')[i].click()
		q_right = '正确答案为：' + q_right.get_attribute('textContent').strip()

		document.add_paragraph(q_right).paragraph_format.line_spacing = Pt(22)
		# print(q_right)
		# print(q_right.text)
		# 答案解析
		# print('答案解析：')
		q_tit = browser.find_elements_by_xpath('//div[@class="d-txt clearfix"]/table/tbody/tr/td')[i].get_attribute('textContent')
		q_tit = '答案解析：\n' + q_tit.encode('GBK','ignore').decode('GBk')
		q_tit = re.sub('(翻译：)', r'\n\1\n', q_tit)
		q_tit = re.sub('(解题思路：)', r'\n\1\n', q_tit)
		q_tit = re.sub('[\u4E00-\u9FA5]([A|B|C|D]、)[\u4E00-\u9FA5]', r'\n\1', q_tit)
		q_tit = re.sub('([A|B|C|D])\.', r'\n\1.', q_tit)
		q_tit = re.sub('(\d)\. ', r'\n\1. ', q_tit)
		# print(q_tit)
		
		document.add_paragraph(q_tit).paragraph_format.line_spacing = Pt(22)
		# 答案解析中也会有图片
		if is_element_exist(img_xpath):
			img_id = 'f-check-answer_' + q_id
			img_xpath = '//div[@id="%s"]//div[@class="d-txt clearfix"]/table/tbody/tr/td//img' % img_id
			img = browser.find_elements_by_xpath(img_xpath)
			for k,im in enumerate(img):
				imgname = q_id + '_tit-%d.png' % k
				img_path = os.path.join(file_path, imgname)
				# print('tit图片地址为：%s' % (img_path, ))
				img_src = im.get_attribute('src')
				try:
					if not os.path.exists(img_path):
						img = requests.get(img_src)
						with open(img_path, 'wb') as f:
							f.write(img.content)
					pic = document.add_picture(img_path)
					last_paragraph = document.paragraphs[-1]
					last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
				except:
					pass
		
	browser.quit()
document.save('new1.docx')
# os.remove('gdCMAe-Part2-%d.docx' % pid_set)
t = time.time()-t0
m, s = divmod(t, 60)
h, m = divmod(m, 60)
print("花费时间：%d:%02d:%02d" % (h, m, s))

